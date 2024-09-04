import streamlit as st
import os
from docx import Document
from PyPDF2 import PdfReader
import pandas as pd
from main import generate_documentation
import logging
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Function to generate prompt
def generate_prompt(text_input=None, file_content=None, template=None):
    template_text = f"Using the {template} template: " if template else ""
    if text_input:
        prompt = template_text + text_input
    elif file_content:
        prompt = template_text + file_content
    else:
        prompt = template_text + "No input provided."
    return prompt

# Function to create a .docx file
def create_docx(text, filename):
    document = Document()
    document.add_paragraph(text)
    document_path = f"{filename}.docx"
    document.save(document_path)
    return document_path

# Function to read file content based on file type
def read_file(uploaded_file):
    if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(uploaded_file)
        file_content = "\n".join([para.text for para in doc.paragraphs])
    elif uploaded_file.type == "application/pdf":
        pdf = PdfReader(uploaded_file)
        file_content = ""
        for page in pdf.pages:
            file_content += page.extract_text()
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        df = pd.read_excel(uploaded_file)
        file_content = df.to_string()
    elif uploaded_file.type == "text/csv":
        df = pd.read_csv(uploaded_file)
        file_content = df.to_string()
    else:
        file_content = None
    return file_content

def main():
    st.title("AI-Powered Documentation Generator")
    st.write("Generate professional documentation effortlessly from text input or uploaded files using OpenAI's advanced language model.")
    
    # CSS for styling
    st.markdown("""
    <style>
    .reportview-container {
        background: #f0f0f5;
        font-family: 'Arial', sans-serif;
    }
    .sidebar .sidebar-content {
        background: #f0f0f5;
    }
    .sidebar .sidebar-content h1 {
        color: #2c3e50;
    }
    .stButton>button {
        background-color: #2c3e50;
        color: white;
    }
    </style>
    """, unsafe_allow_html=True)
    
    st.sidebar.title("Input Parameters")
    method = st.sidebar.radio("Choose Input Method", ("Text Input", "File Upload"))
    template = st.sidebar.selectbox("Choose a Template", ["Standard Report", "Meeting Minutes", "Technical Documentation"])
    
    text_input = ""
    file_content = ""

    if method == "Text Input":
        text_input = st.sidebar.text_area("Enter your text here", "Type your documentation content...", height=300)
    elif method == "File Upload":
        uploaded_file = st.sidebar.file_uploader("Upload your file", type=["docx", "pdf", "xlsx", "csv"])
        if uploaded_file:
            file_content = read_file(uploaded_file)
            if file_content is None:
                st.sidebar.error("Unsupported file type. Please upload a .docx, .pdf, .xlsx, or .csv file.")
                return

    if st.sidebar.button("Generate Documentation"):
        with st.spinner("Generating your documentation..."):
            try:
                prompt = generate_prompt(text_input, file_content, template)
                result = generate_documentation(prompt)
            except Exception as e:
                logger.error(f"Error during documentation process: {e}")
                st.error(f"Error: {e}")
                return

            try:
                docx_path = create_docx(result, "Documentation")
            except Exception as e:
                logger.error(f"Error creating the .docx file: {e}")
                st.error(f"Error: {e}")
                return

            with open(docx_path, "rb") as f:
                btn = st.download_button(
                    label="Download Documentation",
                    data=f,
                    file_name="Documentation.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            st.success("Documentation generated successfully!")
            st.markdown(result)

    if st.sidebar.button("Preview Document"):
        st.markdown(result)

if __name__ == "__main__":
    main()
